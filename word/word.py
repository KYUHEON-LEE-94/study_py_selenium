from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
# 현재 날짜와 시간을 가져옵니다.
now = datetime.now()

doc = Document("example.docx")

current_year = now.year
current_month = now.month
target_company_name = input("기관명 입력: ")
filename = f"{current_year}년 {current_month}월 {target_company_name} 검색서비스 정기점검 보고서.docx"


def create_word_doc():
    doc = Document("example.docx")  # 워드 파일 열기
    
    first_page(doc)
    second_page(doc)

    doc.save(filename)  # 수정된 문서 저장

def set_font(runs, font_name=None, size=None):
    for run in runs:
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if size:
            run.font.size = Pt(size)

def set_text_and_font(para, text, font_name=None, size=None):
    para.text = text
    set_font(para.runs, font_name, size)
    doc.save("modified_example.docx")  # 수정된 문서 저장

def find_paragraph_by_text(doc, target_text):
    for para in doc.paragraphs:
        if para.text == target_text:
            return para
    return None

def first_page(doc):
    for idx, para in enumerate(doc.paragraphs, start=1):
        if idx == 1:
            set_font(para.runs, size=32)
        elif idx == 2:
            set_text_and_font(para, text=f"{target_company_name} 검색 서비스", font_name='맑은 고딕', size=36)
        elif idx == 3:
            set_text_and_font(para, text=f"<{current_month}월 정기점검 보고서>", font_name='맑은 고딕', size=20)


def find_paragraph_by_text(doc, target_text):
    for idx, para in enumerate(doc.paragraphs):
        if para.text == target_text:
            return para, idx
    return None, -1  # 해당하는 문단이 없을 경우 None과 -1을 반환합니다.

def find_paragraph_after_index(doc, start_index):
    for idx, para in enumerate(doc.paragraphs):
        if idx == start_index + 1:
            return para
    return None

def second_page(doc):
    target_text = "검색서비스 정기점검 보고서"
    found_para, idx = find_paragraph_by_text(doc, target_text)
    if found_para:        
        next_para = find_paragraph_after_index(doc, idx)
        if next_para:
            # 문단을 찾았을 경우에 대한 처리
            print("주어진 문단 다음의 문단:", next_para.text)
            # 필요한 작업 수행
        else:
            # 문단을 찾지 못한 경우에 대한 처리
            print("주어진 문단 다음의 문단을 찾을 수 없습니다.")
    else:
        # 문단을 찾지 못한 경우에 대한 처리
        print("문단을 찾을 수 없습니다.")
 

